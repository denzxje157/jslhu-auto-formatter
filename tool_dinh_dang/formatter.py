import docx
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import copy
import io
import os
import sys
import re
import shutil

def set_cell_margins(cell, top=80, bottom=80, left=80, right=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def style_metadata_table(table):
    """
    Đóng khung bảng Metadata (Table 0 và Table 1) trang 1 bằng các đường kẻ ngang
    chuẩn màu đen (Hex: 000000) giống như bài báo mẫu, đồng thời làm sạch viền dọc và padding.
    """
    # Đặt kiểu bảng về Normal Table để xóa sạch mọi viền mặc định của Word
    table.style = 'Normal Table'
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
        
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>' # 1.5 pt top border
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>' # 1.5 pt bottom border
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    
    # 1. Làm sạch toàn bộ viền ô riêng biệt (tcBorders) và cấu hình khoảng cách dòng cực kỳ sát
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cantSplit = trPr.find(qn('w:cantSplit'))
        if cantSplit is None:
            trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            
        for cell in row.cells:
            # Thu nhỏ padding (top/bottom = 2pt / 40 dxa, left/right = 4pt / 80 dxa) để không bị nhảy trang
            set_cell_margins(cell, top=40, bottom=40, left=80, right=80)
            
            # Xóa sạch border ô kế thừa cũ
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = tcPr.find(qn('w:tcBorders'))
            if tcBorders is not None:
                tcPr.remove(tcBorders)
                
            # Đưa giãn dòng về 1.0 và space before/after về 0
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0

    # 2. Tạo đường kẻ ngang phân cách tiêu đề cột (Row 0 - dưới THÔNG TIN BÀI BÁO / ARTICLE INFO)
    if len(table.rows) > 0:
        for cell in table.rows[0].cells[:2]:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)
            
    # 3. Tạo đường kẻ ngang trên hàng TỪ KHÓA / KEYWORDS (Row 5)
    if len(table.rows) > 5:
        for cell in table.rows[5].cells[:2]:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    # 4. Tạo đường kẻ ngang trên & dưới hàng nội dung Từ khóa (Row 6) chuẩn 0.5pt (sz="4")
    if len(table.rows) > 6:
        for c_idx, cell in enumerate(table.rows[6].cells):
            tcPr = cell._tc.get_or_add_tcPr()
            if c_idx < 2:
                tcBorders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    f'</w:tcBorders>'
                )
            else:
                tcBorders = parse_xml(
                    f'<w:tcBorders {nsdecls("w")}>'
                    f'  <w:top w:val="nil"/>'
                    f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    f'</w:tcBorders>'
                )
            tcPr.append(tcBorders)

    # 5. Đặt độ rộng cột cố định để tránh ngắt dòng ngày tháng tiếng Anh (Tổng 7.0 inches)
    col_widths = [Inches(1.2), Inches(1.2), Inches(4.6)]
    for row in table.rows:
        if len(row.cells) == 3:
            row.cells[0].width = col_widths[0]
            row.cells[1].width = col_widths[1]
            row.cells[2].width = col_widths[2]
        elif len(row.cells) == 2:
            row.cells[0].width = col_widths[0] + col_widths[1]
            row.cells[1].width = col_widths[2]

def style_table(table):
    # Căn giữa bảng và tự động kích hoạt tính năng co giãn cột
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = True
    
    # Thiết lập độ rộng cột đều nhau để bảng tự động khớp vừa khít với độ rộng cột báo 3.25 inches
    num_cols = len(table.columns)
    if num_cols > 0:
        col_width = Inches(3.25 / num_cols)
        for row in table.rows:
            for cell in row.cells:
                cell.width = col_width

    # Xóa border cũ
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
        
    # Thiết lập border 3 dòng chuẩn học thuật LHU (3-line table)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    
    # Thêm dòng kẻ dưới cho hàng đầu tiên (Header Row)
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)
            
    # Cấu hình từng ô trong bảng
    for r_idx, row in enumerate(table.rows):
        # Không cắt hàng qua trang
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        # Lặp lại tiêu đề hàng ở đầu mỗi trang
        if r_idx == 0:
            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
            
        for cell in row.cells:
            set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(9)
                    if r_idx == 0:
                        run.bold = True

def replace_text_preserve_formatting(paragraph, new_text, bold=False, italic=False, size_pt=10, align=None, color_rgb=None):
    if not paragraph.runs:
        run = paragraph.add_run(new_text)
    else:
        paragraph.runs[0].text = new_text
        run = paragraph.runs[0]
        for r in paragraph.runs[1:]:
            r.text = ""
            
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb is not None:
        run.font.color.rgb = color_rgb
    if align is not None:
        paragraph.alignment = align

def replace_cell_text(cell, new_text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=9, color_rgb=None):
    if not cell.paragraphs:
        cell.add_paragraph()
    p = cell.paragraphs[0]
    p.alignment = align
    replace_text_preserve_formatting(p, new_text, bold=bold, italic=italic, size_pt=size_pt, color_rgb=color_rgb)
    # Xóa paragraph thừa
    for extra_p in cell.paragraphs[1:]:
        p_elm = extra_p._element
        p_elm.getparent().remove(p_elm)

def fill_merged_abstract_cell(cell, header_text, body_text):
    # Xóa toàn bộ các đoạn văn hiện có trong ô gộp để tránh trùng lặp
    for p in list(cell.paragraphs):
        p_elm = p._element
        p_elm.getparent().remove(p_elm)
        
    # Tạo Paragraph 0 chứa nhãn tiêu đề (Ví dụ: TÓM TẮT)
    p_hdr = cell.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_hdr.paragraph_format.space_before = Pt(0)
    p_hdr.paragraph_format.space_after = Pt(4)
    p_hdr.paragraph_format.line_spacing = 1.0
    run_hdr = p_hdr.add_run(header_text)
    run_hdr.font.name = 'Times New Roman'
    run_hdr.font.size = Pt(9)
    run_hdr.bold = True
    
    # Tạo Paragraph 1 chứa nội dung tóm tắt học thuật thực tế
    p_body = cell.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_body.paragraph_format.space_before = Pt(0)
    p_body.paragraph_format.space_after = Pt(0)
    p_body.paragraph_format.line_spacing = 1.0
    run_body = p_body.add_run(body_text)
    run_body.font.name = 'Times New Roman'
    run_body.font.size = Pt(9)

def append_element_to_body(doc, element):
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        sectPr.addprevious(element)
    else:
        body.append(element)

def update_headers(doc, data):
    """
    Tự động cập nhật Header chạy tiêu đề (Running Header) cho các trang chẵn và lẻ
    loại bỏ các placeholder 'Article name' và 'Author name'.
    """
    authors_en = [a.get("name", "") for a in data.get("authors_en", [])]
    if len(authors_en) == 1:
        author_str = authors_en[0]
    elif len(authors_en) == 2:
        author_str = f"{authors_en[0]}, {authors_en[1]}"
    elif len(authors_en) > 2:
        author_str = f"{authors_en[0]} et al."
    else:
        author_str = "Nguyen Van A et al."
        
    title_en = data.get("title_en", "").strip()
    if len(title_en) > 70:
        title_en_short = title_en[:67] + "..."
    else:
        title_en_short = title_en
        
    for section in doc.sections:
        # Sửa Default Header - Trang lẻ (Article name -> Tiêu đề tiếng Anh rút gọn)
        for p in section.header.paragraphs:
            if "Article name" in p.text:
                p.text = ""
                run = p.add_run(title_en_short)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(8.5)
                run.italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
        # Sửa Even Page Header - Trang chẵn (Author name -> Tên tác giả tiếng Anh)
        if section.even_page_header:
            for p in section.even_page_header.paragraphs:
                if "Author name" in p.text:
                    p.text = ""
                    run = p.add_run(author_str)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(8.5)
                    run.italic = True
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def format_equation_paragraph(p, math_idx):
    p_fmt = p.paragraph_format
    p_fmt.left_indent = Cm(0)
    p_fmt.first_line_indent = Cm(0)
    p_fmt.space_before = Pt(3)
    p_fmt.space_after = Pt(3)
    p_fmt.alignment = None
    
    p_fmt.tab_stops.clear_all()
        
    p_fmt.tab_stops.add_tab_stop(Cm(4.25), docx.enum.text.WD_TAB_ALIGNMENT.CENTER)
    p_fmt.tab_stops.add_tab_stop(Cm(8.5), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
    
    if not p._element.xpath('.//*[local-name()="oMath"]'):
        return
        
    # Ép cỡ chữ của tất cả các run công thức m:r về 10pt (w:sz val="20")
    for mr in p._element.xpath('.//*[local-name()="r" and namespace-uri()="http://schemas.openxmlformats.org/officeDocument/2006/math"]'):
        rPr = mr.xpath('*[local-name()="rPr"]')
        if rPr:
            sz = rPr[0].xpath('*[local-name()="sz"]')
            if sz:
                sz[0].set(qn('w:val'), '20')
            else:
                sz_xml = parse_xml(f'<w:sz {nsdecls("w")} w:val="20"/>')
                rPr[0].insert(0, sz_xml)
        else:
            rPr_xml = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/><w:sz w:val="20"/></w:rPr>')
            mr.insert(0, rPr_xml)
            
    # Thực hiện deepcopy sau khi các phần tử đã được định dạng 10pt thành công
    math_elms = [copy.deepcopy(x) for x in p._element.xpath('.//*[local-name()="oMath"]')]
    
    pPr = p._element.pPr
    p._element.clear()
    if pPr is not None:
        p._element.append(pPr)
        
    tab1_run = parse_xml(f'<w:r {nsdecls("w")}><w:tab/></w:r>')
    p._element.append(tab1_run)
    
    for idx, math_elm in enumerate(math_elms):
        if idx > 0:
            br_run = parse_xml(f'<w:r {nsdecls("w")}><w:br/><w:tab/></w:r>')
            p._element.append(br_run)
        p._element.append(math_elm)
        
    tab2_run = parse_xml(f'<w:r {nsdecls("w")}><w:tab/><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/><w:sz w:val="20"/></w:rPr><w:t>({math_idx})</w:t></w:r>')
    p._element.append(tab2_run)

def change_reference_numbering_to_brackets(p):
    import re
    # 1. Nếu bắt đầu bằng số dạng "1. " hoặc "12. "
    text_combined = ""
    run_indices = []
    for idx, run in enumerate(p.runs):
        text_combined += run.text
        run_indices.append(idx)
        m = re.match(r'^(\d+)\.\s*', text_combined)
        if m:
            num = m.group(1)
            new_prefix = f"[{num}] "
            matched_len = len(m.group(0))
            remaining = text_combined[matched_len:]
            
            # Cập nhật run đầu tiên chứa tiền tố mới
            p.runs[run_indices[0]].text = new_prefix + remaining
            
            # Xóa chữ ở các run trung gian
            for r_idx in run_indices[1:]:
                p.runs[r_idx].text = ""
            break
            
    # 2. Xóa sạch gạch chân và màu xanh của đầu số ngoặc vuông [X]
    prefix_text = ""
    for run in p.runs:
        prefix_text += run.text
        run.underline = False
        run.font.color.rgb = RGBColor(0, 0, 0)
        if re.match(r'^\[\d+\]\s*', prefix_text):
            break

def bold_year_in_reference(p):
    import re
    import copy
    import docx
    from docx.oxml import OxmlElement
    year_pattern = re.compile(r'\((\d{4})\)')
    runs_list = list(p.runs)
    for run in runs_list:
        text = run.text
        match = year_pattern.search(text)
        if match:
            year_str = match.group(1)
            start_idx = match.start()
            end_idx = match.end()
            before_text = text[:start_idx]
            after_text = text[end_idx:]
            
            run.text = before_text + "("
            
            r_year = OxmlElement('w:r')
            if run._r.rPr is not None:
                r_year.append(copy.deepcopy(run._r.rPr))
            
            # Sử dụng Run của python-docx để thiết lập in đậm chuẩn xác
            year_run_obj = docx.text.run.Run(r_year, p)
            year_run_obj.bold = True
            
            t_year = OxmlElement('w:t')
            t_year.text = year_str
            r_year.append(t_year)
            
            run._r.addnext(r_year)
            
            r_after = OxmlElement('w:r')
            if run._r.rPr is not None:
                r_after.append(copy.deepcopy(run._r.rPr))
            t_after = OxmlElement('w:t')
            t_after.text = ")" + after_text
            r_after.append(t_after)
            
            r_year.addnext(r_after)

def set_authors_paragraph(p, authors):
    p.text = ""
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    
    for idx, a in enumerate(authors):
        name = a.get("name", "")
        aff = a.get("affiliation_number", "")
        is_corr = a.get("is_corresponding", False)
        
        r_name = p.add_run(name)
        r_name.font.name = 'Times New Roman'
        r_name.font.size = Pt(12)
        r_name.bold = False
        
        if aff:
            r_aff = p.add_run(aff)
            r_aff.font.name = 'Times New Roman'
            r_aff.font.size = Pt(12)
            r_aff.bold = False
            r_aff.font.superscript = True
            
        if is_corr:
            r_corr = p.add_run("*")
            r_corr.font.name = 'Times New Roman'
            r_corr.font.size = Pt(12)
            r_corr.bold = False
            r_corr.font.superscript = True
            
        if idx < len(authors) - 1:
            r_sep = p.add_run(", ")
            r_sep.font.name = 'Times New Roman'
            r_sep.font.size = Pt(12)
            r_sep.bold = False

def set_affiliation_paragraph(p, aff_num, aff_name):
    p.text = ""
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    
    if aff_num:
        r_num = p.add_run(aff_num)
        r_num.font.name = 'Times New Roman'
        r_num.font.size = Pt(10)
        r_num.font.superscript = True
        r_num.font.italic = True
        
    r_name = p.add_run(" " + aff_name)
    r_name.font.name = 'Times New Roman'
    r_name.font.size = Pt(10)
    r_name.font.italic = True

def format_document(data, template_path, output_path, source_docx_path=None):
    """
    Tự động hóa định dạng bài báo khoa học dựa trên biểu mẫu JSLHU.
    Bố cục trang tiêu đề đối xứng chuẩn 12_CNTT1019.pdf:
    Tác giả, đơn vị công tác và Email tương tác được đưa ra ngoài và CANH PHẢI (WD_ALIGN_PARAGRAPH.RIGHT) theo đúng mẫu LHU.
    Bảng 0 và Bảng 1 được xóa hàng 0, hiển thị tinh gọn và đóng khung xanh LHU.
    Phần thân bài được định dạng 2 cột, bảo toàn XML công thức toán học và ảnh nhúng.
    """
    # 1. Sao chép template sang vị trí đích
    shutil.copy(template_path, output_path)
    doc = docx.Document(output_path)
    
    # Lấy các metadata chính
    title_vi = data.get("title_vi", "").strip().upper()
    title_en = data.get("title_en", "").strip().upper()
    
    # Màu xanh thương hiệu LHU: Hex 365F91
    lhu_blue = RGBColor(0x36, 0x5F, 0x91)
    
    # Danh sách tác giả
    authors_vi_list = []
    for a in data.get("authors_vi", []):
        name = a.get("name", "")
        aff = a.get("affiliation_number", "")
        corr = "*" if a.get("is_corresponding") else ""
        authors_vi_list.append(f"{name}{aff}{corr}")
    authors_vi_str = ", ".join(authors_vi_list)
    
    authors_en_list = []
    for a in data.get("authors_en", []):
        name = a.get("name", "")
        aff = a.get("affiliation_number", "")
        corr = "*" if a.get("is_corresponding") else ""
        authors_en_list.append(f"{name}{aff}{corr}")
    authors_en_str = ", ".join(authors_en_list)
    
    # --- PHẦN 1: ĐIỀN THÔNG TIN TRANG TIÊU ĐỀ (CANH PHẢI TÁC GIẢ & ĐƠN VỊ CÔNG TÁC) ---
    # Đầu tiên lưu trữ tất cả các đối tượng paragraph ban đầu của template trước khi bị dịch chuyển index
    p_title_vi = doc.paragraphs[1]
    p_guide_vi = doc.paragraphs[2]
    p_authors_vi = doc.paragraphs[3]
    p_space = doc.paragraphs[4] # Đoạn trống ngăn cách Table 0 và tiêu đề tiếng Anh
    p_title_en = doc.paragraphs[5]
    p_guide_en = doc.paragraphs[6]
    p_authors_en = doc.paragraphs[7]
    p_aff_en_1 = doc.paragraphs[8]
    p_aff_en_2 = doc.paragraphs[9]
    
    # Tìm đoạn văn chứa section break (đoạn văn cuối cùng của Section 0)
    p_break = None
    break_idx = -1
    for i, p in enumerate(doc.paragraphs):
        pPr = p._p.pPr
        if pPr is not None:
            sectPr = pPr.find(qn('w:sectPr'))
            if sectPr is not None:
                p_break = p
                break_idx = i
                break

    # Cập nhật Tiêu đề & Tác giả tiếng Việt
    replace_text_preserve_formatting(p_title_vi, title_vi or "TIÊU ĐỀ BÀI BÁO TIẾNG VIỆT", bold=True, size_pt=15, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=lhu_blue)
    p_title_vi.paragraph_format.space_before = Pt(0)
    p_title_vi.paragraph_format.space_after = Pt(6)
    
    if data.get("authors_vi"):
        set_authors_paragraph(p_authors_vi, data.get("authors_vi"))
    else:
        replace_text_preserve_formatting(p_authors_vi, "Tác giả 1, Tác giả 2", bold=False, size_pt=12, align=WD_ALIGN_PARAGRAPH.RIGHT)
        p_authors_vi.paragraph_format.right_indent = Cm(0)
        p_authors_vi.paragraph_format.left_indent = Cm(0)
    
    # Cập nhật Tiêu đề & Tác giả tiếng Anh
    replace_text_preserve_formatting(p_title_en, title_en or "TITLE OF THE PAPER IN ENGLISH", bold=True, size_pt=15, align=WD_ALIGN_PARAGRAPH.CENTER, color_rgb=lhu_blue)
    p_title_en.paragraph_format.space_before = Pt(12)
    p_title_en.paragraph_format.space_after = Pt(6)
    
    if data.get("authors_en"):
        set_authors_paragraph(p_authors_en, data.get("authors_en"))
    else:
        replace_text_preserve_formatting(p_authors_en, "Author 1, Author 2", bold=False, size_pt=12, align=WD_ALIGN_PARAGRAPH.RIGHT)
        p_authors_en.paragraph_format.right_indent = Cm(0)
        p_authors_en.paragraph_format.left_indent = Cm(0)

    # Cập nhật Đơn vị công tác tiếng Anh
    affs_en = data.get("affiliations_en", [])
    if len(affs_en) > 0:
        set_affiliation_paragraph(p_aff_en_1, affs_en[0].get("number", ""), affs_en[0].get("name", ""))
    else:
        p_aff_en_1.text = ""
        p_aff_en_1.paragraph_format.space_before = Pt(0)
        p_aff_en_1.paragraph_format.space_after = Pt(0)
        
    if len(affs_en) > 1:
        set_affiliation_paragraph(p_aff_en_2, affs_en[1].get("number", ""), affs_en[1].get("name", ""))
    else:
        p_aff_en_2.text = ""
        p_aff_en_2.paragraph_format.space_before = Pt(0)
        p_aff_en_2.paragraph_format.space_after = Pt(0)

    # Xóa sạch các đoạn hướng dẫn rác (Kiểu chữ và Font type) để dọn diện tích trang
    p_guide_vi._element.getparent().remove(p_guide_vi._element)
    p_guide_en._element.getparent().remove(p_guide_en._element)

    # Thu nhỏ khoảng cách đoạn trống trung gian p_space để tối ưu không gian trang 1
    p_space.text = ""
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(0)
    p_space.paragraph_format.line_spacing = 1.0

    # Chèn Tác giả liên hệ tiếng Anh (*Corresponding Author) phía TRÊN bảng 1
    p_corr_en = doc.add_paragraph()
    p_corr_en.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_corr_en.paragraph_format.right_indent = Cm(0)
    p_corr_en.paragraph_format.left_indent = Cm(0)
    p_corr_en.paragraph_format.space_before = Pt(0)
    p_corr_en.paragraph_format.space_after = Pt(4)
    p_corr_en.paragraph_format.line_spacing = 1.0
    
    r_ast_en = p_corr_en.add_run("*")
    r_ast_en.font.name = 'Times New Roman'
    r_ast_en.font.size = Pt(10)
    r_ast_en.font.superscript = True
    r_ast_en.font.italic = True
    
    r_text_en = p_corr_en.add_run(" Corresponding Author: " + data.get('email_contact', 'nguyenvana@lhu.edu.vn'))
    r_text_en.font.name = 'Times New Roman'
    r_text_en.font.size = Pt(10)
    r_text_en.font.italic = True
    doc.tables[1]._tbl.addprevious(p_corr_en._p)

    # Chèn Đơn vị công tác & Tác giả liên hệ tiếng Việt phía TRÊN bảng 0
    p_aff_vi = doc.add_paragraph()
    p_aff_vi.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_aff_vi.paragraph_format.right_indent = Cm(0)
    p_aff_vi.paragraph_format.left_indent = Cm(0)
    p_aff_vi.paragraph_format.space_before = Pt(0)
    p_aff_vi.paragraph_format.space_after = Pt(4)
    p_aff_vi.paragraph_format.line_spacing = 1.0
    
    affs_vi = data.get("affiliations_vi", [])
    for idx, aff in enumerate(affs_vi):
        num = aff.get("number", "")
        name = aff.get("name", "")
        if num:
            r_num = p_aff_vi.add_run(num)
            r_num.font.name = 'Times New Roman'
            r_num.font.size = Pt(10)
            r_num.font.superscript = True
            r_num.font.italic = True
        r_name = p_aff_vi.add_run(" " + name)
        r_name.font.name = 'Times New Roman'
        r_name.font.size = Pt(10)
        r_name.font.italic = True
        if idx < len(affs_vi) - 1:
            p_aff_vi.add_run("\n")
            
    if data.get("email_contact"):
        p_aff_vi.add_run("\n")
        r_ast_vi = p_aff_vi.add_run("*")
        r_ast_vi.font.name = 'Times New Roman'
        r_ast_vi.font.size = Pt(10)
        r_ast_vi.font.superscript = True
        r_ast_vi.font.italic = True
        
        r_email = p_aff_vi.add_run(" Tác giả liên hệ: " + data.get("email_contact"))
        r_email.font.name = 'Times New Roman'
        r_email.font.size = Pt(10)
        r_email.font.italic = True
        
    doc.tables[0]._tbl.addprevious(p_aff_vi._p)

    # Ghi mã DOI và Available online ở phía DƯỚI bảng 1 (trực tiếp tại đoạn break_p chứa section break)
    if p_break is not None:
        doi_online = f"Doi: https://doi.org/10.16159/jslhu.26.1019\nAvailable online at: https://lhj.vn"
        replace_text_preserve_formatting(p_break, doi_online, size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        p_break.paragraph_format.space_before = Pt(6)
        p_break.paragraph_format.space_after = Pt(6)
        p_break.paragraph_format.line_spacing = 1.0

    # --- BƯỚC 2: ĐIỀN BẢNG THÔNG TIN TÓM TẮT/ABSTRACT (TABLE 0 & TABLE 1) ---
    # Bảng 0: Tiếng Việt
    if len(doc.tables) > 0:
        table_vi = doc.tables[0]
        # Xóa hàng 0 chứa đơn vị công tác cũ để giống bài báo mẫu
        table_vi.rows[0]._tr.getparent().remove(table_vi.rows[0]._tr)
        
        # Điền tiêu đề cột (màu đen) và nội dung tóm tắt qua hàm ghi ô gộp chuyên dụng
        replace_cell_text(table_vi.rows[0].cells[0], "THÔNG TIN BÀI BÁO", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=9)
        fill_merged_abstract_cell(table_vi.rows[0].cells[2], "TÓM TẮT", data.get("abstract_vi", ""))
        
        # Định dạng từ khóa dọc phân tách bởi dấu chấm phẩy
        kw_vi_items = []
        kw_list = data.get("keywords_vi", [])
        for idx, kw in enumerate(kw_list):
            if idx == len(kw_list) - 1:
                kw_vi_items.append(f"{kw}.")
            else:
                kw_vi_items.append(f"{kw};")
        kw_vi_str = "\n".join(kw_vi_items)
        replace_cell_text(table_vi.rows[6].cells[0], kw_vi_str, size_pt=9)
        
        # Các nhãn ngày tháng màu đen
        replace_cell_text(table_vi.rows[1].cells[0], "Ngày nhận:", bold=False, size_pt=9)
        replace_cell_text(table_vi.rows[2].cells[0], "Ngày hoàn thiện:", bold=False, size_pt=9)
        replace_cell_text(table_vi.rows[3].cells[0], "Ngày chấp nhận:", bold=False, size_pt=9)
        replace_cell_text(table_vi.rows[4].cells[0], "Ngày đăng:", bold=False, size_pt=9)
        replace_cell_text(table_vi.rows[5].cells[0], "TỪ KHÓA", bold=True, size_pt=9)
        
        replace_cell_text(table_vi.rows[1].cells[1], "05/11/2025", size_pt=9)
        replace_cell_text(table_vi.rows[2].cells[1], "02/03/2026", size_pt=9)
        replace_cell_text(table_vi.rows[3].cells[1], "02/03/2026", size_pt=9)
        replace_cell_text(table_vi.rows[4].cells[1], "15/03/2026", size_pt=9)
        
        style_metadata_table(table_vi)

    # Bảng 1: Tiếng Anh
    if len(doc.tables) > 1:
        table_en = doc.tables[1]
        # Xóa hàng 0 chứa email cũ
        table_en.rows[0]._tr.getparent().remove(table_en.rows[0]._tr)
        
        # Xóa hàng cuối cùng rỗng (original Row 8, nay là Row 7) để tối ưu không gian đứng
        if len(table_en.rows) > 7:
            table_en.rows[7]._tr.getparent().remove(table_en.rows[7]._tr)
            
        # Điền tiêu đề cột (màu đen) và nội dung abstract trực tiếp
        replace_cell_text(table_en.rows[0].cells[0], "ARTICLE INFO", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=9)
        replace_cell_text(table_en.rows[0].cells[2], "ABSTRACT", bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=9)
        replace_cell_text(table_en.rows[1].cells[2], data.get("abstract_en", ""), align=WD_ALIGN_PARAGRAPH.JUSTIFY, size_pt=9)
        
        # Định dạng keywords tiếng Anh dọc phân tách bởi dấu chấm phẩy
        kw_en_items = []
        kw_list_en = data.get("keywords_en", [])
        for idx, kw in enumerate(kw_list_en):
            if idx == len(kw_list_en) - 1:
                kw_en_items.append(f"{kw}.")
            else:
                kw_en_items.append(f"{kw};")
        kw_en_str = "\n".join(kw_en_items)
        replace_cell_text(table_en.rows[6].cells[0], kw_en_str, size_pt=9)
        
        # Các nhãn tiếng Anh màu đen
        replace_cell_text(table_en.rows[1].cells[0], "Received:", bold=False, size_pt=9)
        replace_cell_text(table_en.rows[2].cells[0], "Revised:", bold=False, size_pt=9)
        replace_cell_text(table_en.rows[3].cells[0], "Accepted:", bold=False, size_pt=9)
        replace_cell_text(table_en.rows[4].cells[0], "Published:", bold=False, size_pt=9)
        replace_cell_text(table_en.rows[5].cells[0], "KEYWORDS", bold=True, size_pt=9)
        
        replace_cell_text(table_en.rows[1].cells[1], "Nov 5th, 2025", size_pt=9)
        replace_cell_text(table_en.rows[2].cells[1], "Mar 2nd, 2026", size_pt=9)
        replace_cell_text(table_en.rows[3].cells[1], "Mar 2nd, 2026", size_pt=9)
        replace_cell_text(table_en.rows[4].cells[1], "Mar 15th, 2026", size_pt=9)
        
        style_metadata_table(table_en)

    # --- BƯỚC 3: DỌN DẸP PHẦN THÂN CŨ TRONG TEMPLATE ---
    # Tìm vị trí thực tế hiện tại của p_break trong doc.paragraphs bằng cách đối chiếu XML element
    current_break_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if p._element is p_break._element:
            current_break_idx = idx
            break
            
    if current_break_idx != -1:
        # Sử dụng vòng lặp while an toàn để tránh bỏ sót paragraph khi danh sách bị co lại động
        while len(doc.paragraphs) > current_break_idx + 1:
            p_to_del = doc.paragraphs[current_break_idx + 1]
            p_elm = p_to_del._element
            p_elm.getparent().remove(p_elm)
        
    # Xóa các bảng thừa trong mẫu từ index 2 trở đi
    t_len = len(doc.tables)
    for idx in range(t_len - 1, 1, -1):
        t_elm = doc.tables[idx]._element
        t_elm.getparent().remove(t_elm)

    # --- BƯỚC 4: CẤU HÌNH SECTION 2 CÓ SẴN TRONG TEMPLATE ---
    for sec in doc.sections:
        sec.top_margin = Cm(1.5)
        sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(2.5)
        sec.right_margin = Cm(1.0)
        body_section = doc.sections[1]
        body_section.start_type = docx.enum.section.WD_SECTION_START.NEW_PAGE
        
        sectPr = body_section._sectPr
        cols = sectPr.xpath('w:cols')
        if cols:
            # Xóa các thẻ con w:col ghi đè khoảng cách cột
            for child in list(cols[0]):
                cols[0].remove(child)
            cols[0].set(qn('w:num'), '2')
            cols[0].set(qn('w:space'), '288')
            cols[0].set(qn('w:equalWidth'), '1')
        else:
            cols_xml = parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="288" w:equalWidth="1"/>')
            sectPr.append(cols_xml)

    # --- BƯỚC 5: SAO CHÉP & ĐỊNH DẠNG NỘI DUNG THÂN BÀI ---
    if source_docx_path and os.path.exists(source_docx_path):
        with open(source_docx_path, 'rb') as f:
            src_doc = docx.Document(f)
            
        copy_started = False
        ref_started = False
        
        def iter_blocks(parent_doc):
            for child in parent_doc.element.body:
                if child.tag.endswith('p'):
                    yield docx.text.paragraph.Paragraph(child, parent_doc)
                elif child.tag.endswith('tbl'):
                    yield docx.table.Table(child, parent_doc)
                    
        math_counter = 0
        for item in iter_blocks(src_doc):
            if isinstance(item, docx.text.paragraph.Paragraph):
                text = item.text.strip()
                
                # Nhận diện điểm bắt đầu của thân bài (Chương 1)
                if not copy_started:
                    if re.match(r'^(1\.|I\.)\s', text) or "giới thiệu" in text.lower() or "introduction" in text.lower():
                        copy_started = True
                    else:
                        continue
                        
                # Nhận diện phần Tài liệu tham khảo
                if "tài liệu tham khảo" in text.lower() or "references" in text.lower():
                    ref_started = True
                    p_ref_title = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
                    p_ref_title_run = p_ref_title.add_run("5. Tài liệu tham khảo")
                    p_ref_title_run.font.name = 'Times New Roman'
                    p_ref_title_run.font.size = Pt(10)
                    p_ref_title_run.bold = True
                    p_ref_title_run.font.color.rgb = lhu_blue
                    p_ref_title.paragraph_format.space_before = Pt(6)
                    p_ref_title.paragraph_format.space_after = Pt(6)
                    p_ref_title.paragraph_format.left_indent = Inches(0) # Không thụt lề đầu đề TLTK
                    p_ref_title.paragraph_format.first_line_indent = Inches(0)
                    append_element_to_body(doc, p_ref_title._p)
                    continue
                
                # --- NÂNG CẤP CHUẨN HÓA TIÊU ĐỀ THÔNG MINH ---
                is_heading = False
                if re.match(r'^(\d+|[I|V|X|L]+)(\.\d+)*\.?\s+[A-ZĐ]', text) or item.style.name.startswith('Heading'):
                    # Đảm bảo đây không phải là đoạn văn chứa nội dung bài viết dài (như "4.1. Kết luận: Nghiên cứu này...")
                    # Tiêu đề chuẩn phải ngắn hơn 120 ký tự và không chứa dấu hai chấm ':' ngăn cách thân bài
                    if len(text) < 120 and ":" not in text:
                        is_heading = True
                else:
                    low_text = text.lower()
                    if (low_text.startswith("giới thiệu") or 
                        low_text.startswith("phương pháp") or 
                        low_text.startswith("thiết kế") or 
                        low_text.startswith("nội dung") or 
                        low_text.startswith("kết quả") or 
                        low_text.startswith("thực nghiệm") or 
                        low_text.startswith("thảo luận") or 
                        low_text.startswith("kết luận")):
                        if len(text) < 100 and not text.endswith('.') and ":" not in text:
                            is_heading = True

                # Nếu là tiêu đề chính thức, tự động chuẩn hóa văn bản
                if is_heading:
                    normalized_text = text
                    if "giới thiệu" in text.lower() or "introduction" in text.lower():
                        if not re.match(r'^\d+\.', text):
                            normalized_text = "1. " + text
                    elif "phương pháp" in text.lower() or "thiết kế" in text.lower() or "kiến trúc" in text.lower():
                        if not re.match(r'^\d+\.', text):
                            normalized_text = "2. " + text
                    elif "kết quả" in text.lower() or "thực nghiệm" in text.lower() or "thảo luận" in text.lower():
                        if not re.match(r'^\d+\.', text):
                            normalized_text = "3. " + text
                    elif "kết luận" in text.lower() or "conclusion" in text.lower():
                        if not re.match(r'^\d+\.', text):
                            normalized_text = "4. " + text
                            
                    p_copy = copy.deepcopy(item._p)
                    append_element_to_body(doc, p_copy)
                    new_p = docx.text.paragraph.Paragraph(p_copy, doc)
                    replace_text_preserve_formatting(new_p, normalized_text, bold=True, size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT, color_rgb=lhu_blue)
                else:
                    p_copy = copy.deepcopy(item._p)
                    append_element_to_body(doc, p_copy)
                    new_p = docx.text.paragraph.Paragraph(p_copy, doc)
                    if ref_started:
                        change_reference_numbering_to_brackets(new_p)
                
                # Định dạng font chữ của đoạn văn vừa copy
                for run in new_p.runs:
                    run.font.name = 'Times New Roman'
                    if is_heading:
                        run.font.size = Pt(10)
                        run.bold = True
                        run.font.color.rgb = lhu_blue
                    elif ref_started:
                        run.font.size = Pt(9)
                        if re.match(r'^\d{4}$', run.text.strip()):
                            run.bold = True
                        else:
                            run.bold = False
                    else:
                        run.font.size = Pt(10)
                        # Giữ nguyên định dạng in đậm/in thường gốc của từng từ trong thân bài
                        
                # Định dạng kiểu lề và khoảng cách đoạn văn
                p_fmt = new_p.paragraph_format
                if is_heading:
                    p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_fmt.left_indent = Inches(0) # KHÔNG THỤT VÀO lề trái theo đúng bài đăng mẫu (Hình 3)
                    p_fmt.first_line_indent = Inches(0) # KHÔNG THỤT VÀO dòng đầu
                    p_fmt.space_before = Pt(6)
                    p_fmt.space_after = Pt(6)
                    p_fmt.line_spacing = 1.0
                elif ref_started:
                    p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_fmt.left_indent = Cm(0.5)
                    p_fmt.first_line_indent = -Cm(0.5)
                    p_fmt.space_before = Pt(0)
                    p_fmt.space_after = Pt(4)
                    p_fmt.line_spacing = 1.0
                else:
                    # Chú thích hình/bảng
                    if text.startswith("Hình ") or text.startswith("Bảng ") or text.lower().startswith("figure ") or text.lower().startswith("table "):
                        p_fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_fmt.left_indent = Inches(0)
                        p_fmt.first_line_indent = Inches(0)
                        p_fmt.space_before = Pt(4)
                        p_fmt.space_after = Pt(4)
                        for r in new_p.runs:
                            r.font.size = Pt(9)
                            r.italic = True
                            r.bold = True
                    # Tiêu đề công thức (nếu chỉ chứa text chữ)
                    elif "Công thức" in text and len(new_p._element.xpath('.//*[local-name()="oMath"]')) == 0:
                        p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_fmt.first_line_indent = Cm(0.36)
                        p_fmt.left_indent = Cm(0)
                        p_fmt.space_before = Pt(0)
                        p_fmt.space_after = Pt(6)
                        p_fmt.line_spacing = 1.0
                        for r in new_p.runs:
                            r.font.size = Pt(10)
                            r.bold = False
                            r.italic = False
                    # Bản thân công thức toán học (chứa m:oMath)
                    elif len(new_p._element.xpath('.//*[local-name()="oMath"]')) > 0:
                        math_counter += 1
                        format_equation_paragraph(new_p, math_counter)
                    else:
                        p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_fmt.first_line_indent = Cm(0.36) # Thụt dòng đầu 0.36 cm
                        p_fmt.left_indent = Cm(0)
                        p_fmt.space_before = Pt(0)
                        p_fmt.space_after = Pt(6)
                        p_fmt.line_spacing = 1.0
                        
                # Xử lý ảnh nhúng
                drawings = item._p.xpath('.//w:drawing')
                if drawings:
                    for dr in new_p._p.xpath('.//w:drawing'):
                        dr.getparent().remove(dr)
                    for dr in drawings:
                        embed_ids = dr.xpath('.//a:blip/@r:embed')
                        if embed_ids:
                            rId = embed_ids[0]
                            try:
                                image_part = src_doc.part.related_parts[rId]
                                image_bytes = image_part.blob
                                image_stream = io.BytesIO(image_bytes)
                                
                                run = new_p.add_run()
                                run.add_picture(image_stream, width=Inches(3.25))
                                new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p_fmt.first_line_indent = Inches(0)
                            except Exception as e:
                                print(f"Error copying image: {e}")
                                
            elif isinstance(item, docx.table.Table):
                if not copy_started:
                    continue
                tbl_copy = copy.deepcopy(item._tbl)
                append_element_to_body(doc, tbl_copy)
                new_table = docx.table.Table(tbl_copy, doc)
                style_table(new_table)
                
    else:
        # CHẾ ĐỘ FALLBACK
        for sec in data.get("sections", []):
            title = sec.get("title", "")
            content = sec.get("content", "")
            
            p_title = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
            p_title_run = p_title.add_run(title)
            p_title_run.font.name = 'Times New Roman'
            p_title_run.font.size = Pt(10)
            p_title_run.bold = True
            p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_title.paragraph_format.left_indent = Inches(0) # Sửa lùi lề về 0
            p_title.paragraph_format.first_line_indent = Inches(0)
            p_title.paragraph_format.space_before = Pt(6)
            p_title.paragraph_format.space_after = Pt(6)
            p_title.paragraph_format.line_spacing = 1.0
            append_element_to_body(doc, p_title._p)
            
            paragraphs_text = content.split('\n')
            for p_text in paragraphs_text:
                if not p_text.strip():
                    continue
                p_body = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
                p_body_run = p_body.add_run(p_text.strip())
                p_body_run.font.name = 'Times New Roman'
                p_body_run.font.size = Pt(10)
                p_body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_body.paragraph_format.first_line_indent = Inches(0.14)
                p_body.paragraph_format.space_before = Pt(0)
                p_body.paragraph_format.space_after = Pt(0)
                p_body.paragraph_format.line_spacing = 1.0
                append_element_to_body(doc, p_body._p)
                
        ack = data.get("acknowledgments", "")
        if ack:
            p_ack_title = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
            p_ack_title_run = p_ack_title.add_run("Lời cảm ơn")
            p_ack_title_run.font.name = 'Times New Roman'
            p_ack_title_run.font.size = Pt(10)
            p_ack_title_run.bold = True
            p_ack_title.paragraph_format.left_indent = Inches(0)
            p_ack_title.paragraph_format.space_before = Pt(6)
            append_element_to_body(doc, p_ack_title._p)
            
            p_ack = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
            p_ack_run = p_ack.add_run(ack)
            p_ack_run.font.name = 'Times New Roman'
            p_ack_run.font.size = Pt(10)
            p_ack.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_ack.paragraph_format.first_line_indent = Inches(0.14)
            append_element_to_body(doc, p_ack._p)
            
        refs = data.get("references", [])
        if refs:
            p_ref_title = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
            p_ref_title_run = p_ref_title.add_run("5. Tài liệu tham khảo")
            p_ref_title_run.font.name = 'Times New Roman'
            p_ref_title_run.font.size = Pt(10)
            p_ref_title_run.bold = True
            p_ref_title.paragraph_format.left_indent = Inches(0)
            p_ref_title.paragraph_format.space_before = Pt(6)
            append_element_to_body(doc, p_ref_title._p)
            
            for r in refs:
                p_ref = docx.text.paragraph.Paragraph(OxmlElement('w:p'), doc)
                p_ref_run = p_ref.add_run(r)
                p_ref_run.font.name = 'Times New Roman'
                p_ref_run.font.size = Pt(9)
                change_reference_numbering_to_brackets(p_ref)
                p_ref.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_ref.paragraph_format.left_indent = Cm(0.5)
                p_ref.paragraph_format.first_line_indent = -Cm(0.5)
                p_ref.paragraph_format.space_after = Pt(4)
                p_ref.paragraph_format.line_spacing = 1.0
                append_element_to_body(doc, p_ref._p)
                
    # --- BƯỚC 6: CẬP NHẬT HEADER CHẠY CỦA TÀI LIỆU ---
    update_headers(doc, data)
                
    doc.save(output_path)
    print("Format document complete.")
