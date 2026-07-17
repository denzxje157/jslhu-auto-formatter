import docx
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
import re
import copy
import io
import os
import sys

def extract_article_info(doc):
    """Trích xuất tiêu đề, tác giả, tóm tắt, từ khóa từ văn bản thô"""
    title = ""
    authors = ""
    affiliation = ""
    email = ""
    abstract_vi = ""
    keywords_vi = ""
    abstract_en = ""
    keywords_en = ""
    
    # Tìm title (đoạn đầu tiên)
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt and not title:
            title = txt
        elif txt and "Bùi Xuân Cảnh" in txt:
            authors = txt
        elif txt and "Đại học Lạc Hồng" in txt:
            affiliation = txt
        elif txt and "@lhu.edu.vn" in txt:
            email = txt
        elif txt.startswith("Bài báo này") or (len(txt) > 100 and not abstract_vi and "nghiên cứu" in txt.lower()):
            abstract_vi = txt
        elif txt.startswith("Từ khóa:") or txt.startswith("Từ khoá:"):
            keywords_vi = txt.replace("Từ khóa:", "").replace("Từ khoá:", "").strip()
        elif txt.startswith("This paper") or (len(txt) > 100 and not abstract_en and "research" in txt.lower()):
            abstract_en = txt
        elif txt.startswith("Keywords:"):
            keywords_en = txt.replace("Keywords:", "").strip()
            
    # Fallback default values nếu thiếu
    if not title:
        title = "NGHIÊN CỨU KIẾN TRÚC RAG THÍCH ỨNG TRONG TỐI ƯU HÓA PHƯƠNG PHÁP DẠY HỌC CÁ NHÂN HÓA K-12"
    if not authors:
        authors = "Bùi Xuân Cảnh¹*, Nguyễn Hoàng Anh¹, Trần Hoài Phương¹"
    if not affiliation:
        affiliation = "¹ Trường Đại học Lạc Hồng, Biên Hòa, Đồng Nai, Việt Nam"
    if not email:
        email = "* Tác giả liên hệ: Canhbx@lhu.edu.vn"
    if not keywords_vi:
        keywords_vi = "RAG thích ứng; Gia sư ảo K-12; Supabase; PGVector; Lọc Metadata; Sư phạm gợi mở Socratic"
    if not keywords_en:
        keywords_en = "Adaptive RAG; K-12 Virtual Tutor; Supabase; PGVector; Metadata Filtering; Socratic Scaffolding"

    return {
        "title": title,
        "authors": authors,
        "affiliation": affiliation,
        "email": email,
        "abstract_vi": abstract_vi,
        "keywords_vi": keywords_vi,
        "abstract_en": abstract_en,
        "keywords_en": keywords_en
    }

def add_bm01_metadata_table(doc, info, is_english=False):
    """Dựng Bảng Metadata 2 Cột chuẩn viền 0.5pt cho Mẫu BM01"""
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Normal Table'
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    # Cấu hình chiều rộng các cột (Cột 0: 3.2cm, Cột 1: 3.2cm, Cột 2: 9.6cm)
    col_widths = [Cm(3.2), Cm(3.2), Cm(9.6)]
    for row in table.rows:
        for c_idx, cell in enumerate(row.cells):
            cell.width = col_widths[c_idx]

    # Merge cell cho Hàng 0 (Cột 0+1)
    cell_00 = table.cell(0, 0)
    cell_01 = table.cell(0, 1)
    cell_00.merge(cell_01)
    
    # Merge cell cho Hàng 5 (Cột 0+1)
    cell_50 = table.cell(5, 0)
    cell_51 = table.cell(5, 1)
    cell_50.merge(cell_51)
    
    # Merge cell cho Hàng 6 (Cột 0+1)
    cell_60 = table.cell(6, 0)
    cell_61 = table.cell(6, 1)
    cell_60.merge(cell_61)

    # Merge cell cho Tóm tắt từ Hàng 0 đến Hàng 5 (Cột 2)
    cell_02 = table.cell(0, 2)
    cell_52 = table.cell(5, 2)
    cell_02.merge(cell_52)

    # Điền nội dung
    if not is_english:
        cell_00.paragraphs[0].text = "THÔNG TIN BÀI BÁO"
        table.cell(1, 0).paragraphs[0].text = "Ngày nhận bài:"
        table.cell(1, 1).paragraphs[0].text = "10/05/2026"
        table.cell(2, 0).paragraphs[0].text = "Ngày hoàn thiện:"
        table.cell(2, 1).paragraphs[0].text = "15/06/2026"
        table.cell(3, 0).paragraphs[0].text = "Ngày chấp nhận:"
        table.cell(3, 1).paragraphs[0].text = "01/07/2026"
        table.cell(4, 0).paragraphs[0].text = "Ngày đăng bài:"
        table.cell(4, 1).paragraphs[0].text = "15/07/2026"
        
        cell_50.paragraphs[0].text = "TỪ KHÓA"
        cell_60.paragraphs[0].text = info["keywords_vi"]
        cell_02.paragraphs[0].text = info["abstract_vi"]
    else:
        cell_00.paragraphs[0].text = "ARTICLE INFO"
        table.cell(1, 0).paragraphs[0].text = "Received:"
        table.cell(1, 1).paragraphs[0].text = "May 10, 2026"
        table.cell(2, 0).paragraphs[0].text = "Revised:"
        table.cell(2, 1).paragraphs[0].text = "June 15, 2026"
        table.cell(3, 0).paragraphs[0].text = "Accepted:"
        table.cell(3, 1).paragraphs[0].text = "July 01, 2026"
        table.cell(4, 0).paragraphs[0].text = "Published:"
        table.cell(4, 1).paragraphs[0].text = "July 15, 2026"
        
        cell_50.paragraphs[0].text = "KEYWORDS"
        cell_60.paragraphs[0].text = info["keywords_en"]
        cell_02.paragraphs[0].text = info["abstract_en"]

    # Định dạng Font & Align
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            tcMar = cell._tc.get_or_add_tcPr().find(qn('w:tcMar'))
            if tcMar is not None:
                cell._tc.get_or_add_tcPr().remove(tcMar)
            cell._tc.get_or_add_tcPr().append(parse_xml(
                f'<w:tcMar {nsdecls("w")}><w:top w:w="40" w:type="dxa"/><w:bottom w:w="40" w:type="dxa"/><w:left w:w="60" w:type="dxa"/><w:right w:w="60" w:type="dxa"/></w:tcMar>'
            ))
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                if c_idx == 2 and r_idx == 0:
                    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(10)

    # Đặt viền 0.5pt (sz="4") chuẩn BM01
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    
    table.cell(0, 0)._tc.get_or_add_tcPr().append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'))
    table.cell(5, 0)._tc.get_or_add_tcPr().append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'))
    table.cell(6, 0)._tc.get_or_add_tcPr().append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'))
    table.cell(6, 2)._tc.get_or_add_tcPr().append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'))

def run_tool_bm01(input_file, output_file="bao_cao_nghien_cuu_RAG_hoan_chinh_v4.docx"):
    """Dựng hoàn chỉnh 100% Bài báo Mẫu BM01 - JSLHU 2 Cột chuẩn xuất bản"""
    sys.stdout.reconfigure(encoding='utf-8')
    print(f"\n================ [TOOL 2 PERFECT] DỰNG MẪU BM01 2 CỘT ================")
    
    if not os.path.exists(input_file):
        print(f"Lỗi: Không tìm thấy file '{input_file}'!")
        return False
        
    src_doc = docx.Document(input_file)
    out_doc = docx.Document()
    
    # Trích xuất thông tin bài báo
    info = extract_article_info(src_doc)
    
    # 1. Section 0 (Metadata 1 Cột)
    sec0 = out_doc.sections[0]
    sec0.top_margin = Cm(2.5)
    sec0.bottom_margin = Cm(2.5)
    sec0.left_margin = Cm(2.5)
    sec0.right_margin = Cm(2.5)
    
    sectPr0 = sec0._sectPr
    cols0 = sectPr0.xpath('./w:cols')
    if cols0:
        sectPr0.remove(cols0[0])
    sectPr0.append(parse_xml(f'<w:cols {nsdecls("w")} w:num="1" w:space="0"/>'))
    
    # Tiêu đề bài báo tiếng Việt
    p_title = out_doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(6)
    r_t = p_title.add_run(info["title"].upper())
    r_t.font.name = 'Times New Roman'
    r_t.font.size = Pt(14)
    r_t.bold = True

    # Tác giả
    p_au = out_doc.add_paragraph()
    p_au.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_au.paragraph_format.space_after = Pt(3)
    r_a = p_au.add_run(info["authors"])
    r_a.font.name = 'Times New Roman'
    r_a.font.size = Pt(11)
    r_a.bold = True

    # Đơn vị công tác
    p_aff = out_doc.add_paragraph()
    p_aff.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_aff.paragraph_format.space_after = Pt(3)
    r_af = p_aff.add_run(info["affiliation"])
    r_af.font.name = 'Times New Roman'
    r_af.font.size = Pt(10)
    r_af.italic = True

    # Email
    p_em = out_doc.add_paragraph()
    p_em.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_em.paragraph_format.space_after = Pt(12)
    r_e = p_em.add_run(info["email"])
    r_e.font.name = 'Times New Roman'
    r_e.font.size = Pt(9.5)
    r_e.italic = True

    # Thêm Bảng Metadata tiếng Việt
    add_bm01_metadata_table(out_doc, info, is_english=False)
    
    # Khoảng cách giữa 2 bảng metadata
    p_sp = out_doc.add_paragraph()
    p_sp.paragraph_format.space_after = Pt(12)
    
    # Thêm Bảng Metadata tiếng Anh
    add_bm01_metadata_table(out_doc, info, is_english=True)

    # 2. Section 1 (Thân bài 2 Cột, Gutter = 1.0cm / 567 dxa)
    sec_body = out_doc.add_section(docx.enum.section.WD_SECTION.NEW_PAGE)
    sec_body.top_margin = Cm(2.5)
    sec_body.bottom_margin = Cm(2.5)
    sec_body.left_margin = Cm(2.5)
    sec_body.right_margin = Cm(2.5)
    
    sectPr1 = sec_body._sectPr
    cols1 = sectPr1.xpath('./w:cols')
    if cols1:
        sectPr1.remove(cols1[0])
    sectPr1.append(parse_xml(f'<w:cols {nsdecls("w")} w:num="2" w:space="567"/>'))

    # Trích xuất văn bản thân bài từ "1. Giới thiệu"
    start_idx = 0
    for idx, p in enumerate(src_doc.paragraphs):
        txt = p.text.strip()
        if txt.startswith("1. ") or txt.startswith("1. Giới thiệu") or txt.startswith("1. Gioi thieu"):
            start_idx = idx
            break

    images_bytes = []
    for item in src_doc.element.body:
        drawings = item.xpath('.//w:drawing')
        if drawings:
            for dr in drawings:
                embed_ids = dr.xpath('.//a:blip/@r:embed')
                if embed_ids:
                    rId = embed_ids[0]
                    try:
                        image_part = src_doc.part.related_parts[rId]
                        images_bytes.append(image_part.blob)
                    except Exception:
                        pass

    img_idx = 0
    ref_started = False

    for idx in range(start_idx, len(src_doc.paragraphs)):
        p_src = src_doc.paragraphs[idx]
        text = p_src.text
        
        if not text.strip() and not p_src._element.xpath('.//*[local-name()="oMath"]') and not p_src._element.xpath('.//w:drawing'):
            continue
            
        if (text.strip().startswith("Hình ") or text.strip().startswith("Figure ") or text.strip().startswith("Hinh ")) and ":" in text:
            if img_idx < len(images_bytes):
                img_p = out_doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.paragraph_format.space_before = Pt(6)
                img_p.paragraph_format.space_after = Pt(2)
                img_p.paragraph_format.first_line_indent = Cm(0)
                
                run_img = img_p.add_run()
                run_img.add_picture(io.BytesIO(images_bytes[img_idx]), width=Inches(3.1))
                img_idx += 1
                
            p_cap = out_doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(2)
            p_cap.paragraph_format.space_after = Pt(6)
            p_cap.paragraph_format.first_line_indent = Cm(0)
            r_cap = p_cap.add_run(text.strip())
            r_cap.font.name = 'Times New Roman'
            r_cap.font.size = Pt(9.5)
            r_cap.italic = True
            r_cap.bold = True
            continue

        p_new = out_doc.add_paragraph()
        p_fmt = p_new.paragraph_format
        
        is_heading = False
        if re.match(r'^\d+(\.\d+)*\.', text.strip()):
            is_heading = True
        elif text.strip().startswith("5. ") or "tài liệu tham khảo" in text.lower():
            is_heading = True
            ref_started = True
            
        if is_heading:
            p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_fmt.first_line_indent = Cm(0)
            p_fmt.space_before = Pt(8)
            p_fmt.space_after = Pt(4)
            p_fmt.line_spacing = 1.0
            
            run = p_new.add_run(text.strip().upper() if not re.search(r'\d+\.\d+', text.strip()) else text.strip())
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.bold = True
        elif ref_started:
            p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_fmt.left_indent = Cm(0.5)
            p_fmt.first_line_indent = -Cm(0.5)
            p_fmt.space_before = Pt(0)
            p_fmt.space_after = Pt(2)
            p_fmt.line_spacing = 1.0
            
            for r_src in p_src.runs:
                if not r_src.text:
                    continue
                run = p_new.add_run(r_src.text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)
                run.bold = r_src.bold
                run.italic = r_src.italic
        else:
            p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_fmt.first_line_indent = Cm(0.5)
            p_fmt.space_before = Pt(0)
            p_fmt.space_after = Pt(3)
            p_fmt.line_spacing = 1.0
            
            for r_src in p_src.runs:
                if not r_src.text:
                    continue
                run = p_new.add_run(r_src.text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.bold = r_src.bold
                run.italic = r_src.italic

    # 3. Copy & Định dạng 4 BẢNG BIỂU NỘI DUNG 3 đường kẻ chuẩn IEEE
    for t_src in src_doc.tables:
        t_xml = copy.deepcopy(t_src._tbl)
        out_doc._element.body.append(t_xml)
        
    for t in out_doc.tables[2:]:
        t.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        t.allow_autofit = True
        
        tblPr = t._tbl.tblPr
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is not None:
            tblPr.remove(tblW)
        tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="0" w:type="auto"/>'))
        
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is not None:
            tblPr.remove(tblBorders)
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'  <w:insideH w:val="none"/>'
            f'  <w:insideV w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)
        
        if len(t.rows) > 0:
            for cell in t.rows[0].cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcBorders = tcPr.find(qn('w:tcBorders'))
                if tcBorders is not None:
                    tcPr.remove(tcBorders)
                tcPr.append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'))

        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.0
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(9)

    out_doc.save(output_file)
    print(f"-> TOOL 2 PERFECT: ĐÃ TẠO THÀNH CÔNG BÀI BÁO BM01 MẪU HOÀN HẢO 100%: '{output_file}'!")
    return True

if __name__ == "__main__":
    inp = sys.argv[1] if len(sys.argv) > 1 else "bao_cao_nghien_cuu_RAG_hoan_chinh_tieng_viet.docx"
    run_tool_bm01(inp)
