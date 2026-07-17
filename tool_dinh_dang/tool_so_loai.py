import docx
from docx.shared import Cm, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
import re
import copy
import io
import os
import sys

def run_tool_so_loai(input_file, output_file="output_ban_nop_so_loai.docx"):
    """
    TOOL 1: Định dạng Vòng Sơ Loại Online Portal (1 Cột A4, Top 3cm, Bottom 2cm, Left 3cm, Right 2cm, Font 11pt, after 3pt)
    Duyệt tuần tự văn bản để giữ nguyên 100% bố cục inline của Hình ảnh, Bảng biểu và Văn bản thô.
    """
    sys.stdout.reconfigure(encoding='utf-8')
    print(f"\n================ [TOOL 1] BẮT ĐẦU ĐỊNH DẠNG VÒNG SƠ LOẠI ================")
    print(f"File đầu vào: '{input_file}'")
    
    if not os.path.exists(input_file):
        print(f"Lỗi: Không tìm thấy file '{input_file}'!")
        return False
        
    src_doc = docx.Document(input_file)
    out_doc = docx.Document()
    
    # 1. Cấu hình Page Setup 1 Cột A4 chuẩn sơ loại
    section = out_doc.sections[0]
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.2)
    
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if cols:
        sectPr.remove(cols[0])
    sectPr.append(parse_xml(f'<w:cols {nsdecls("w")} w:num="1" w:space="0"/>'))

    # 2. Trích xuất tất cả Hình ảnh nhúng từ file nguồn
    images_bytes = []
    for item in src_doc.element.body:
        drawings = item.xpath('.//w:drawing')
        if drawings:
            for dr in drawings:
                embed_ids = dr.xpath('.//a:blip/@r:embed')
                if embed_ids:
                    rId = embed_ids[0]
                    try:
                        image_part = src_doc.part.related_parts[rId]
                        images_bytes.append(image_part.blob)
                    except Exception:
                        pass
                        
    img_idx = 0
    
    # 3. Duyệt tuần tự các phần tử của body để bảo vệ thứ tự tuyệt đối
    for child in src_doc.element.body:
        tag = child.tag.split('}')[-1]
        
        if tag == 'p':
            p_src = docx.text.paragraph.Paragraph(child, src_doc)
            text = p_src.text
            
            # Bỏ qua các đoạn trống không chứa công thức hay hình ảnh
            if not text.strip() and not p_src._element.xpath('.//*[local-name()="oMath"]') and not p_src._element.xpath('.//w:drawing'):
                continue
                
            # Nhận diện tiêu đề Hình ảnh để chèn hình ngay phía trước
            if (text.strip().startswith("Hình ") or text.strip().startswith("Figure ") or text.strip().startswith("Hinh ")) and ":" in text:
                if img_idx < len(images_bytes):
                    img_p = out_doc.add_paragraph()
                    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    img_p.paragraph_format.space_before = Pt(6)
                    img_p.paragraph_format.space_after = Pt(3)
                    img_p.paragraph_format.first_line_indent = Cm(0)
                    
                    run_img = img_p.add_run()
                    run_img.add_picture(io.BytesIO(images_bytes[img_idx]), width=Inches(5.8)) # Căn chuẩn 5.8 inches
                    img_idx += 1
                    
                # Thêm caption cho hình
                p_cap = out_doc.add_paragraph()
                p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_cap.paragraph_format.space_before = Pt(3)
                p_cap.paragraph_format.space_after = Pt(6)
                p_cap.paragraph_format.first_line_indent = Cm(0)
                r_cap = p_cap.add_run(text.strip())
                r_cap.font.name = 'Times New Roman'
                r_cap.font.size = Pt(10)
                r_cap.italic = True
                r_cap.bold = True
                continue
                
            p_new = out_doc.add_paragraph()
            p_fmt = p_new.paragraph_format
            
            # Heading tiêu đề hoặc Tóm tắt/Abstract
            is_heading = False
            if re.match(r'^\d+(\.\d+)*\.', text.strip()):
                is_heading = True
            elif text.strip().startswith("TÓM TẮT") or text.strip().startswith("ABSTRACT") or text.strip().startswith("Từ khóa:") or text.strip().startswith("Keywords:"):
                is_heading = True
                
            if is_heading:
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_fmt.first_line_indent = Cm(0)
                p_fmt.space_before = Pt(6)
                p_fmt.space_after = Pt(6)
                p_fmt.line_spacing = 1.0
                
                run = p_new.add_run(text.strip())
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.bold = True
            else:
                p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_fmt.first_line_indent = Cm(0.5)
                p_fmt.space_before = Pt(0)
                p_fmt.space_after = Pt(3)
                p_fmt.line_spacing = 1.0
                
                for r_src in p_src.runs:
                    if not r_src.text:
                        continue
                    run = p_new.add_run(r_src.text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.bold = r_src.bold
                    run.italic = r_src.italic
                    
        elif tag == 'tbl':
            t_src = docx.table.Table(child, src_doc)
            t_xml = copy.deepcopy(t_src._tbl)
            out_doc._element.body.append(t_xml)
            
            # Định dạng bảng vừa chèn chuẩn 3 đường kẻ
            t = out_doc.tables[-1]
            t.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
            t.allow_autofit = True
            
            tblPr = t._tbl.tblPr
            tblW = tblPr.find(qn('w:tblW'))
            if tblW is not None:
                tblPr.remove(tblW)
            tblPr.append(parse_xml(f'<w:tblW {nsdecls("w")} w:w="0" w:type="auto"/>'))
            
            tblBorders = tblPr.find(qn('w:tblBorders'))
            if tblBorders is not None:
                tblPr.remove(tblBorders)
            borders = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="000000"/>'
                f'  <w:left w:val="none"/>'
                f'  <w:right w:val="none"/>'
                f'  <w:insideH w:val="none"/>'
                f'  <w:insideV w:val="none"/>'
            	f'</w:tblBorders>'
            )
            tblPr.append(borders)
            
            if len(t.rows) > 0:
                for cell in t.rows[0].cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)
                    tcPr.append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tcBorders>'))

            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        p.paragraph_format.space_before = Pt(2)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.line_spacing = 1.0
                        for r in p.runs:
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(9.5)

    out_doc.save(output_file)
    print(f"-> TOOL 1 hoàn thành xuất file: '{output_file}'!")
    return True

if __name__ == "__main__":
    inp = sys.argv[1] if len(sys.argv) > 1 else "bao_cao_nghien_cuu_RAG_ban_nop_so_loai.docx"
    run_tool_so_loai(inp)
